import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional
import hashlib
from datetime import datetime

@dataclass
class ParsedDocument:
    doc_id: str
    filename: str
    source_path: str
    file_type: str
    title: Optional[str] = None
    author: Optional[str] = None
    creation_date: Optional[str] = None
    pages: List[str] = field(default_factory=list)
    tables: List[str] = field(default_factory=list)
    raw_metadata: dict = field(default_factory=dict)

class DocumentParser:
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.txt'}
    
    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = path.suffix.lower()
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {ext}")
        
        doc_id = self._generate_doc_id(file_path)
        
        if ext == '.pdf':
            return self._parse_pdf(file_path, doc_id, path.name)
        elif ext == '.docx':
            return self._parse_docx(file_path, doc_id, path.name)
        elif ext == '.txt':
            return self._parse_txt(file_path, doc_id, path.name)
    
    def _parse_pdf(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        # Primary extraction with PyMuPDF
        doc = fitz.open(file_path)
        metadata = doc.metadata
        pages = []
        
        for page in doc:
            text = page.get_text("text")
            pages.append(text)
        
        doc.close()
        
        # Extract tables with pdfplumber
        tables = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_tables = page.extract_tables()
                for table in page_tables:
                    if table:
                        table_str = self._table_to_markdown(table)
                        tables.append(table_str)
        
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            source_path=file_path,
            file_type='pdf',
            title=metadata.get('title', None),
            author=metadata.get('author', None),
            creation_date=metadata.get('creationDate', None),
            pages=pages,
            tables=tables,
            raw_metadata=dict(metadata)
        )
    
    def _parse_docx(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        doc = DocxDocument(file_path)
        
        # Extract metadata
        props = doc.core_properties
        pages = []
        current_page = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                current_page.append(para.text)
            elif current_page:
                pages.append("\n".join(current_page))
                current_page = []
        
        if current_page:
            pages.append("\n".join(current_page))
        
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            source_path=file_path,
            file_type='docx',
            title=props.title,
            author=props.author,
            creation_date=str(props.created) if props.created else None,
            pages=pages,
            raw_metadata={
                "title": props.title,
                "author": props.author,
                "created": str(props.created) if props.created else None
            }
        )
    
    def _parse_txt(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Split into pages by double newline or fixed chunk
        pages = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            source_path=file_path,
            file_type='txt',
            pages=pages
        )
    
    def _table_to_markdown(self, table: List[List]) -> str:
        if not table:
            return ""
        markdown = ""
        # Header
        header = table[0]
        markdown += "| " + " | ".join(str(c) if c else "" for c in header) + " |\n"
        markdown += "| " + " | ".join("---" for _ in header) + " |\n"
        # Rows
        for row in table[1:]:
            markdown += "| " + " | ".join(str(c) if c else "" for c in row) + " |\n"
        return markdown
    
    def _generate_doc_id(self, file_path: str) -> str:
        return hashlib.md5(file_path.encode()).hexdigest()[:12]

    def parse_directory(self, directory_path: str) -> List[ParsedDocument]:
        path = Path(directory_path)
        print("Searching in:", path.resolve())
        documents = []
        for file_path in path.rglob("*"):
            if file_path.suffix.lower() in self.supported_formats:
                try:
                    doc = self.parse(str(file_path))
                    documents.append(doc)
                    print(f"Parsed: {file_path.name}")
                except Exception as e:
                    print(f"Failed: {file_path.name} - {e}")
        return documents